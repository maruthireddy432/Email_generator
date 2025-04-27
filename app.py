import streamlit as st
import os
from langchain_groq import ChatGroq
from langchain_community.document_loaders import WebBaseLoader
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
import fitz
from docx import Document

st.set_page_config(page_title="Cold Email Generator", layout="centered")

st.markdown("""
    <style>
        .main-title {
            text-align: center;
            font-size: 36px;
            font-weight: bold;
            color: #4A90E2;
        }
        .subtitle {
            text-align: center;
            font-size: 20px;
            color: #999;
            margin-bottom: 20px;
        }
        .stButton>button {
            width: 100%;
            background-color: #4A90E2;
            color: white;
            font-size: 18px;
            padding: 10px;
            border-radius: 8px;
        }
    </style>
""", unsafe_allow_html=True)

# Page Title and Description
st.markdown("<div class='main-title'>Cold Email Generator</div>", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Craft a professional and personalized email for recruiters in seconds!</div>", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Provide the below details !</div>", unsafe_allow_html=True)

# Extract API Key securely
def get_api_key():
    """Fetch API key from environment variable or file"""
    key = os.getenv("GROQ_API_KEY")
    if key:
        return key
    files = [file for file in os.listdir('.') if file.endswith('.txt')]
    if not files:
        st.error("No API key found. Please set 'GROQ_API_KEY' environment variable or provide a key file.")
        return None
    with open(files[0], 'r') as file:
        return file.read().strip()

api_key = get_api_key()
if not api_key:
    st.stop()

# Initialize LangChain chatbot
llm = ChatGroq(
    temperature=1,
    groq_api_key=api_key,
    model_name="gemma2-9b-it"
)

resume = st.file_uploader("Upload your resume (PDF or DOCX)", type=['pdf', 'docx'])

def extract_text_from_pdf(file):
    """Extract text from a PDF file"""
    try:
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in pdf:
            text += page.get_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(file):
    """Extract text from a DOCX file"""
    try:
        doc = Document(file)
        text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text.append("\t".join(row_text))  # Join columns with tab spacing

        return "\n".join(text)
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None

# Process resume after upload
resume_text = None
if resume is not None:
    if resume.name.endswith('.pdf'):
        resume_text = extract_text_from_pdf(resume)
    elif resume.name.endswith('.docx'):
        resume_text = extract_text_from_docx(resume)

# User selection for input method
option = st.radio("Select input method:", ("🔗 Provide URL", "📝 Add Description"))

if option == "🔗 Provide URL":
    url = st.text_input("Enter the URL of the career's page")
elif option == "📝 Add Description":
    job_description = st.text_area("Enter the job description")

def extract_text(url):
    """Extract text from the given URL using WebBaseLoader"""
    try:
        loader = WebBaseLoader(url)
        page_data = loader.load()
        return page_data[0].page_content if page_data else None
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return None

def transform_text(text):
    """Extract job details in JSON format"""
    prompt_extract = PromptTemplate.from_template(
        """
        ### SCRAPED TEXT FROM WEBSITE:
        {page_data}
        ### INSTRUCTION:
        Extract job postings from the career page and return them in JSON format with 
        keys: `role`, `experience`, `skills`, `qualification`, and `description`.
        Only return valid JSON.
        ### VALID JSON (NO PREAMBLE):    
        """
    )
    chain_extract = prompt_extract | llm
    res = chain_extract.invoke({'page_data': text})
    return res.content if res else None

def parse_text(text):
    """Parse JSON-formatted job details"""
    try:
        parser = JsonOutputParser()
        return parser.invoke(text)
    except Exception as e:
        st.error(f"Error parsing JSON: {e}")
        return None

def extract_text_resume(text):
    """Extract structured information from a resume"""
    prompt_extract = PromptTemplate.from_template(
        """
        ### DOCUMENT TEXT:
        {text}
        ### INSTRUCTION:
        Extract key details from the resume and return the following in JSON format:
        - Name
        - Email
        - Phone Number
        - Skills
        - Experience
        - Certifications
        - Projects
        - Summary
        Only return valid JSON.
        ### VALID JSON (NO PREAMBLE):    
        """
    )
    res_text = prompt_extract | llm
    res = res_text.invoke({'text': text})
    return parse_text(res.content)

parsed_resume_text = extract_text_resume(resume_text) if resume_text else None

def generate_email(parsed_text, parsed_resume_text):
    """Generate a cold email based on job description and resume"""
    email_prompt = PromptTemplate.from_template(
        """
        ### RESUME DETAILS:
        {text}
        ### JOB DESCRIPTION:
        {description}
        ### INSTRUCTION:
        Write a concise, engaging, and professional cold email to a recruiter.
        - Keep the tone warm and conversational.
        - Highlight relevant skills and experience from the resume.
        - Avoid bullet points.
        - Include a clear call to action.
        - place the relevant information like name and email from resume text in the email.
        - Keep the email under 200 words.
        (NO PREAMBLE)
        """
    )
    chain_email = email_prompt | llm
    email_text = chain_email.invoke({'text': parsed_resume_text, 'description': parsed_text})
    return email_text if email_text else None

if st.button("Generate Email"):
    if option == "🔗 Provide URL":
        if not url:
            st.error("Please enter a valid URL.")
        else:
            text = extract_text(url)
            if not text:
                st.error("Failed to extract text from the provided URL.")
            else:
                parsed_text = parse_text(transform_text(text))
                if parsed_text and parsed_resume_text:
                    email = generate_email(parsed_text, parsed_resume_text)
                    if email:
                        st.subheader("Generated Cold Email:")
                        st.write(email.content)
                    else:
                        st.error("Failed to generate an email.")
    else:
        if not job_description:
            st.error("Please enter a job description.")
        else:
            parsed_text = parse_text(transform_text(job_description))
            if parsed_text and parsed_resume_text:
                email = generate_email(parsed_text, parsed_resume_text)
                if email:
                    st.subheader("Generated Cold Email:")
                    st.write(email.content)
                else:
                    st.error("Failed to generate an email.")
